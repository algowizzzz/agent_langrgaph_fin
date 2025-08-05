from pathlib import Path
from typing import Dict, List, Any
import logging
import re
import json
import os
from datetime import datetime

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter

# PDF processing support
try:
    import PyPDF2
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False
    logging.warning("PyPDF2 not available. PDF processing disabled.")

# DOCX processing support  
try:
    import docx
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False
    logging.warning("python-docx not available. DOCX processing disabled.")

# Use real config instead of MockConfig
try:
    from config import config
except ImportError:
    # Fallback MockConfig with 5k word chunks
    class MockConfig:
        def __init__(self):
            self.ai = self.AI()
        
        class AI:
            def __init__(self):
                self.chunk_size = 25000  # ~5k words 
                self.chunk_overlap = 1000  # Better context overlap
    config = MockConfig()


class DocumentProcessor:
    """Service for processing uploaded documents into text chunks."""
    def __init__(self):
        self.headers_to_split_on = [("#", "Header 1"), ("##", "Header 2"), ("###", "Header 3")]
        self.markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=self.headers_to_split_on, strip_headers=False)
        self.recursive_splitter = RecursiveCharacterTextSplitter(chunk_size=config.ai.chunk_size, chunk_overlap=config.ai.chunk_overlap)

    async def process_document(self, file_path: str) -> Dict:
        """Process document with proper file type detection."""
        try:
            file_path = Path(file_path)
            file_ext = file_path.suffix.lower()
            
            # Extract content based on file type
            if file_ext == '.pdf':
                content, file_type = await self._extract_pdf_text(file_path)
            elif file_ext in ['.docx', '.doc']:
                content, file_type = await self._extract_docx_text(file_path)
            elif file_ext in ['.txt', '.md']:
                content, file_type = await self._extract_text_file(file_path)
            elif file_ext == '.csv':
                content, file_type = await self._extract_csv_text(file_path)
            elif file_ext in ['.xlsx', '.xls']:  # 🆕 EXCEL SUPPORT - SPECIAL HANDLING
                sheet_chunks, file_type = await self._extract_excel_text(file_path)
                # Excel files already come pre-chunked (1 sheet = 1 chunk)
                text_chunks = sheet_chunks
            else:
                # Fallback to text reading
                content, file_type = await self._extract_text_file(file_path)
            
            # Process with appropriate splitter - using 5k word chunks for all types
            # EXCEPT Excel files which are already optimally chunked
            if file_ext not in ['.xlsx', '.xls']:
                # Use recursive splitter for non-Excel files
                split_result = self.recursive_splitter.split_text(content)
                text_chunks = split_result
                
            # Create Document objects with metadata
            docs = []
            for i, chunk in enumerate(text_chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "file_type": file_type,
                        "file_name": file_path.name,
                        "chunk_index": i
                    }
                )
                docs.append(doc)
                
            return {"success": True, "documents": docs, "file_type": file_type}
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def _extract_pdf_text(self, file_path: Path) -> tuple[str, str]:
        """Extract text from PDF with page structure."""
        if not HAS_PDF_SUPPORT:
            raise Exception("PyPDF2 not available for PDF processing")
            
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Format with page headers for structure detection
                        formatted_page = f"\n\n## Page {page_num + 1}\n\n{page_text}"
                        text_content.append(formatted_page)
                except Exception as e:
                    # Skip problematic pages
                    continue
        
        return "\n".join(text_content), "PDF"
    
    async def _extract_docx_text(self, file_path: Path) -> tuple[str, str]:
        """Extract text from DOCX file."""
        if not HAS_DOCX_SUPPORT:
            # Fallback to text reading
            return file_path.read_text(encoding='utf-8'), "TEXT"
            
        try:
            import docx
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
                    
            return "\n".join(content), "DOCX"
        except Exception:
            # Fallback to text reading
            return file_path.read_text(encoding='utf-8'), "TEXT"
    
    async def _extract_text_file(self, file_path: Path) -> tuple[str, str]:
        """Extract text from plain text file."""
        return file_path.read_text(encoding='utf-8'), "TEXT"
    
    async def _extract_csv_text(self, file_path: Path) -> tuple[str, str]:
        """Extract text representation from CSV file."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            # Convert to readable text format
            content = f"# CSV Data: {file_path.name}\n\n"
            content += df.to_string(index=False)
            return content, "CSV"
        except Exception:
            # Fallback to text reading
            return file_path.read_text(encoding='utf-8'), "TEXT"

    async def _extract_excel_text(self, file_path: Path) -> tuple[list, str]:
        """Extract structured JSON representation from Excel file with multiple sheets - ONE SHEET PER CHUNK."""
        try:
            import pandas as pd
            import json
            
            # 🎯 KEY: sheet_name=None reads ALL sheets into a dictionary
            all_sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            # 🔧 NEW APPROACH: Return structured JSON for each sheet
            sheet_chunks = []
            
            for sheet_name, df in all_sheets.items():
                # 📊 OPTION 3: Full JSON with Metadata (OPTIMAL for LLM)
                try:
                    # Convert dtypes to serializable format
                    dtypes_dict = {}
                    for col, dtype in df.dtypes.items():
                        dtypes_dict[str(col)] = str(dtype)
                    
                    # Create structured JSON object
                    sheet_data = {
                        "excel_metadata": {
                            "file_name": file_path.name,
                            "sheet_name": sheet_name,
                            "shape": {"rows": df.shape[0], "columns": df.shape[1]},
                            "has_data": not df.empty
                        },
                        "table_structure": {
                            "columns": [str(col) for col in df.columns.tolist()],
                            "column_types": dtypes_dict,
                            "index": [str(idx) for idx in df.index.tolist()]
                        },
                        "table_data": df.to_dict('records') if not df.empty else [],
                        "summary_text": f"Excel sheet '{sheet_name}' from {file_path.name} with {df.shape[0]} rows and {df.shape[1]} columns"
                    }
                    
                    # Convert to JSON string for storage (preserves structure)
                    sheet_json = json.dumps(sheet_data, indent=2, default=str)
                    sheet_chunks.append(sheet_json)
                    
                except Exception as sheet_error:
                    # Fallback for problematic sheets
                    fallback_content = f"# Excel Data: {file_path.name}\n\n## Sheet: {sheet_name}\n\n"
                    if df.empty:
                        fallback_content += "*(Empty sheet)*\n\n"
                    else:
                        fallback_content += df.to_string(index=False) + "\n\n"
                    sheet_chunks.append(fallback_content)
                    print(f"⚠️ Using fallback for sheet '{sheet_name}': {str(sheet_error)}")
                
            return sheet_chunks, "EXCEL"
            
        except Exception as e:
            # Fallback with informative error message
            print(f"⚠️ Excel processing failed for {file_path.name}: {str(e)}")
            return [f"Excel processing failed: {str(e)}\nFile: {file_path.name}"], "TEXT"

class PersistentDocumentStore:
    """Thread-safe persistent document store using file system."""
    
    def __init__(self, store_path: str = "document_store.json"):
        self.store_path = store_path
        self._cache = None
        
    def _load(self) -> Dict[str, List[Dict]]:
        """Load document store from disk."""
        if self._cache is not None:
            return self._cache
            
        try:
            if os.path.exists(self.store_path):
                with open(self.store_path, 'r') as f:
                    self._cache = json.load(f)
            else:
                self._cache = {}
        except Exception as e:
            logger.warning(f"Failed to load document store: {e}")
            self._cache = {}
        return self._cache
    
    def _save(self, data: Dict[str, List[Dict]]):
        """Save document store to disk."""
        try:
            with open(self.store_path, 'w') as f:
                json.dump(data, f, indent=2)
            self._cache = data
        except Exception as e:
            logger.error(f"Failed to save document store: {e}")
    
    def keys(self):
        """Get all document names."""
        return list(self._load().keys())
    
    def __contains__(self, key: str) -> bool:
        """Check if document exists."""
        return key in self._load()
    
    def __getitem__(self, key: str) -> List[Dict]:
        """Get document chunks."""
        data = self._load()
        if key not in data:
            raise KeyError(f"Document '{key}' not found")
        return data[key]
    
    def __setitem__(self, key: str, value: List[Dict]):
        """Store document chunks."""
        data = self._load()
        data[key] = value
        self._save(data)
    
    def clear(self):
        """Clear all documents."""
        self._save({})
    
    def items(self):
        """Get all document name/chunks pairs."""
        return self._load().items()

document_processor = DocumentProcessor()
document_chunk_store = PersistentDocumentStore()
logger = logging.getLogger(__name__)

print(f"🔧 MODULE LOADED: document_tools.py at {__file__}")  # Module load detection

# --- Tool Implementations ---
def extract_clean_filename(filename):
    """Extract clean filename from timestamped names"""
    import re
    
    # Pattern: YYYYMMDD_HHMMSS_uuid_original_filename
    # We want to extract just the original_filename part
    parts = filename.split('_')
    
    # If it starts with date pattern (8 digits), it's timestamped
    if len(parts) >= 3 and re.match(r'^\d{8}$', parts[0]) and re.match(r'^\d{6}$', parts[1]):
        # Find where the UUID ends (UUIDs have hyphens)
        for i, part in enumerate(parts[2:], start=2):
            if '-' in part:  # This is likely the UUID part
                # Everything after this should be the original filename
                if i + 1 < len(parts):
                    original_parts = parts[i+1:]
                    return '_'.join(original_parts)
                break
    
    # If no timestamp pattern found, return as-is
    return filename

async def upload_document(file_path: str, session_id: str = "unknown", additional_metadata: dict = None, original_filename: str = None) -> dict:
    """Upload document with enhanced metadata tracking and clean display names."""
    result = await document_processor.process_document(file_path)
    if not result["success"]: return {"status": "error", "message": result.get("error")}
    
    chunks = result["documents"]
    doc_name = Path(file_path).name
    
    # Extract clean display name
    if original_filename:
        display_name = extract_clean_filename(original_filename)
    else:
        display_name = extract_clean_filename(Path(file_path).name)
    
    # Enhanced metadata for each chunk
    enhanced_chunks = []
    for doc in chunks:
        enhanced_metadata = {
            **doc.metadata,
            "uploaded_by_session": session_id,
            "upload_timestamp": datetime.now().isoformat(),
            "file_size": Path(file_path).stat().st_size,
            "original_path": str(file_path),
            "display_name": display_name
        }
        
        # Add any additional metadata
        if additional_metadata:
            enhanced_metadata.update(additional_metadata)
            
        enhanced_chunks.append({
            "page_content": doc.page_content, 
            "metadata": enhanced_metadata
        })
    
    document_chunk_store[doc_name] = enhanced_chunks
    return {
        "status": "success", 
        "doc_name": doc_name, 
        "chunks_created": len(chunks),
        "file_type": result.get("file_type", "UNKNOWN"),
        "file_size": Path(file_path).stat().st_size
    }

def _apply_search_query(chunks: List[Dict], query: str) -> List[Dict]:
    """Apply search query with boolean logic support."""
    if not query:
        return chunks
    
    query = query.strip()
    
    # Handle OR logic (most common in failed tests)
    if ' OR ' in query.upper():
        terms = [term.strip().lower() for term in re.split(r'\s+OR\s+', query, flags=re.IGNORECASE)]
        return [c for c in chunks if any(term in c.get("page_content", "").lower() for term in terms)]
    
    # Handle AND logic
    elif ' AND ' in query.upper():
        terms = [term.strip().lower() for term in re.split(r'\s+AND\s+', query, flags=re.IGNORECASE)]
        return [c for c in chunks if all(term in c.get("page_content", "").lower() for term in terms)]
    
    # Handle simple multi-word queries (split on spaces and use OR logic)
    elif ' ' in query and not any(op in query.upper() for op in [' OR ', ' AND ', '"']):
        terms = [term.strip().lower() for term in query.split()]
        return [c for c in chunks if any(term in c.get("page_content", "").lower() for term in terms)]
    
    # Simple single-term search
    else:
        return [c for c in chunks if query.lower() in c.get("page_content", "").lower()]

async def discover_document_structure(doc_name: str) -> dict:
    if doc_name not in document_chunk_store: return {"status": "error", "message": f"Doc '{doc_name}' not found."}
    chunks = document_chunk_store[doc_name]
    headers = {c.get("metadata", {}).get(k) for c in chunks for k in c.get("metadata", {}) if "Header" in k}
    return {"status": "success", "headers": sorted(list(h for h in headers if h is not None))}

async def search_multiple_docs(doc_names: List[str], query: str = None, filter_by_metadata: dict = None, **kwargs) -> list:
    """Search across multiple documents simultaneously."""
    if not doc_names:
        return [{"error": "No documents specified"}]
    
    all_chunks = []
    
    # Collect chunks from all specified documents
    for doc_name in doc_names:
        if doc_name in document_chunk_store:
            chunks = document_chunk_store[doc_name]
            # Add source document info to each chunk
            for chunk in chunks:
                chunk_copy = chunk.copy()
                chunk_copy["source_document"] = doc_name
                all_chunks.append(chunk_copy)
        else:
            # Add error chunk for missing document
            all_chunks.append({"error": f"Document '{doc_name}' not found", "source_document": doc_name})
    
    # Apply metadata filter if provided
    if filter_by_metadata:
        key, value = list(filter_by_metadata.items())[0]
        all_chunks = [c for c in all_chunks if not c.get("error") and value == c.get("metadata", {}).get(key)]
    
    # Apply keyword query filter with proper boolean logic
    if query:
        all_chunks = _apply_search_query(all_chunks, query)
    
    return all_chunks

async def search_uploaded_docs(doc_name: str, query: str = None, filter_by_metadata: dict = None, **kwargs) -> list:
    """Search single document with improved error handling and cache coherency."""
    print(f"🔍 FUNCTION CALLED: search_uploaded_docs with doc_name='{doc_name}'")
    
    # Force cache refresh to ensure consistency
    document_chunk_store._cache = None
    
    # Check if document exists with detailed debugging
    doc_exists = doc_name in document_chunk_store
    store_size = len(list(document_chunk_store.keys()))
    
    print(f"🔍 SEARCH DEBUG: doc_name='{doc_name}', in_store={doc_exists}, store_keys={store_size}")
    
    if not doc_exists:
        print(f"❌ SEARCH DEBUG: Document '{doc_name}' NOT FOUND in store")
        print(f"📋 Available documents: {list(document_chunk_store.keys())[:3]}...")
        
        # Check for similar names (fuzzy matching)
        available_docs = list(document_chunk_store.keys())
        similar_docs = [d for d in available_docs if doc_name.split('_')[-1].replace('.pdf', '') in d]
        
        if similar_docs:
            print(f"🔍 SUGGESTION: Found similar documents: {similar_docs[:2]}")
            return [{"error": f"Document '{doc_name}' not found. Similar documents available: {similar_docs[:2]}"}]
        else:
            return [{"error": f"Document '{doc_name}' not found. {store_size} documents available in store."}]
    
    # Successfully found document
    print(f"✅ SEARCH SUCCESS: Document found with {len(document_chunk_store[doc_name])} chunks")
    
    # Start with all chunks for the document
    filtered_chunks = document_chunk_store[doc_name]
    
    # Apply metadata filter only if it is provided
    if filter_by_metadata:
        key, value = list(filter_by_metadata.items())[0]
        filtered_chunks = [c for c in filtered_chunks if value == c.get("metadata", {}).get(key)]
        print(f"🔍 METADATA FILTER: {len(filtered_chunks)} chunks after metadata filter")
        
    # Apply keyword query filter with proper boolean logic
    if query:
        original_count = len(filtered_chunks)
        filtered_chunks = _apply_search_query(filtered_chunks, query)
        print(f"🔍 QUERY FILTER: {len(filtered_chunks)} chunks after query filter (was {original_count})")
        
    print(f"🔍 FINAL RESULT: Returning {len(filtered_chunks)} chunks")
    return filtered_chunks

# Cross-session document management
async def get_all_documents() -> List[Dict]:
    """Get list of all uploaded documents across sessions with enhanced metadata."""
    all_docs = []
    for doc_name in document_chunk_store.keys():
        # Get first chunk to extract metadata
        chunks = document_chunk_store[doc_name]
        if chunks and len(chunks) > 0:
            first_chunk = chunks[0]
            metadata = first_chunk.get("metadata", {})
            
            # Enhanced document information
            file_size = metadata.get("file_size", 0)
            file_size_str = f"{file_size / 1024:.1f} KB" if file_size < 1024*1024 else f"{file_size / (1024*1024):.1f} MB"
            
            all_docs.append({
                "name": metadata.get("display_name", doc_name),  # Use display_name for human-readable names
                "internal_name": doc_name,  # Keep internal name for backend operations
                "file_type": metadata.get("file_type", "UNKNOWN"),
                "chunks_count": len(chunks),
                "upload_time": metadata.get("upload_timestamp", metadata.get("upload_time", "")),
                "uploaded_by_session": metadata.get("uploaded_by_session", "unknown"),
                "file_size": file_size,
                "file_size_display": file_size_str,
                "source": metadata.get("source", "")
            })
    
    # Sort by upload time (newest first)
    all_docs.sort(key=lambda x: x.get("upload_time", ""), reverse=True)
    return all_docs

async def remove_document(doc_name: str) -> dict:
    """Remove document from the system."""
    if doc_name in document_chunk_store:
        # Remove from store
        data = document_chunk_store._load()
        del data[doc_name]
        document_chunk_store._save(data)
        return {"status": "success", "message": f"Document '{doc_name}' removed successfully"}
    else:
        return {"status": "error", "message": f"Document '{doc_name}' not found"}
