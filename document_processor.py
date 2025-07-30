import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import io
from dataclasses import dataclass
import asyncio

# Document processing imports
import pandas as pd
from docx import Document as DocxDocument
from openpyxl import load_workbook
import PyPDF2
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import config

logger = logging.getLogger(__name__)

@dataclass
class ProcessingResult:
    """Result of document processing."""
    success: bool
    documents: List[Document] = None
    error: Optional[str] = None
    file_type: Optional[str] = None
    
    def __post_init__(self):
        if self.documents is None:
            self.documents = []

class DocumentProcessor:
    """Service for processing uploaded documents into text chunks."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.ai.chunk_size,
            chunk_overlap=config.ai.chunk_overlap,
            length_function=len,
        )
    
    async def process_document(self, file_path: str, timeout_seconds: int = None) -> ProcessingResult:
        """Process a document file into text chunks with timeout handling."""
        timeout = timeout_seconds or config.ai.timeout_seconds
        
        try:
            # Run processing with timeout
            result = await asyncio.wait_for(
                self._process_document_sync(file_path),
                timeout=timeout
            )
            return result
            
        except asyncio.TimeoutError:
            logger.error(f"Document processing timed out after {timeout}s: {file_path}")
            return ProcessingResult(
                success=False,
                error=f"Processing timed out after {timeout} seconds"
            )
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return ProcessingResult(
                success=False,
                error=f"Processing failed: {str(e)}"
            )
    
    async def _process_document_sync(self, file_path: str) -> ProcessingResult:
        """Synchronous document processing (wrapped in async for timeout)."""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return ProcessingResult(
                    success=False,
                    error="File does not exist"
                )
            
            file_ext = file_path.suffix.lower()
            
            # Route to appropriate processor
            if file_ext == '.pdf':
                content = self._process_pdf(file_path)
                file_type = "PDF"
            elif file_ext in ['.doc', '.docx']:
                content = self._process_docx(file_path)
                file_type = "Word Document"
            elif file_ext == '.xlsx':
                content = self._process_xlsx(file_path)
                file_type = "Excel Spreadsheet"
            elif file_ext == '.csv':
                content = self._process_csv(file_path)
                file_type = "CSV File"
            elif file_ext == '.txt':
                content = self._process_txt(file_path)
                file_type = "Text File"
            else:
                return ProcessingResult(
                    success=False,
                    error=f"Unsupported file type: {file_ext}"
                )
            
            if not content or not content.strip():
                return ProcessingResult(
                    success=False,
                    error="No text content could be extracted from the file"
                )
            
            # Split content into chunks
            documents = self.text_splitter.create_documents(
                [content],
                metadatas=[{
                    "source": str(file_path),
                    "file_type": file_type,
                    "file_name": file_path.name
                }]
            )
            
            logger.info(f"Successfully processed {file_path.name}: {len(documents)} chunks, {len(content)} characters")
            
            return ProcessingResult(
                success=True,
                documents=documents,
                file_type=file_type
            )
            
        except Exception as e:
            logger.error(f"Error in document processing: {str(e)}")
            return ProcessingResult(
                success=False,
                error=str(e)
            )
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(cell for cell in row_text):
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- Table ---")
                    text_content.extend(table_text)
                    text_content.append("--- End Table ---")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Word document processing error: {str(e)}")
    
    def _process_xlsx(self, file_path: Path) -> str:
        """Extract text from Excel spreadsheet."""
        try:
            workbook = load_workbook(file_path, read_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                try:
                    sheet = workbook[sheet_name]
                    text_content.append(f"--- Sheet: {sheet_name} ---")
                    
                    # Get data from sheet
                    sheet_data = []
                    for row in sheet.iter_rows(values_only=True):
                        # Filter out completely empty rows
                        if any(cell is not None and str(cell).strip() for cell in row):
                            row_data = [str(cell) if cell is not None else "" for cell in row]
                            sheet_data.append(" | ".join(row_data))
                    
                    if sheet_data:
                        text_content.extend(sheet_data)
                    else:
                        text_content.append("(Empty sheet)")
                    
                except Exception as e:
                    logger.warning(f"Error processing sheet {sheet_name}: {str(e)}")
                    text_content.append(f"Error processing sheet {sheet_name}")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Excel processing error: {str(e)}")
    
    def _process_csv(self, file_path: Path) -> str:
        """Extract structured text from CSV file using pandas for better context."""
        try:
            # Use error_bad_lines=False and warn_bad_lines=True for robustness
            df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='warn')
        except Exception as e:
            logger.warning(f"Initial CSV read failed with utf-8, trying latin-1. Error: {e}")
            try:
                df = pd.read_csv(file_path, encoding='latin-1', on_bad_lines='warn')
            except Exception as final_e:
                raise Exception(f"CSV processing error after multiple encoding attempts: {final_e}")

        # Convert the dataframe to a string format that's easy for an LLM to understand
        # This provides structure without overwhelming the model
        return df.to_string()

    def _process_txt(self, file_path: Path) -> str:
        """Extract text from a plain text file."""
        try:
            return file_path.read_text(encoding='utf-8')
        except Exception as e:
            raise Exception(f"Text file processing error: {str(e)}")
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get statistics about document processing capabilities."""
        return {
            "supported_formats": [".pdf", ".doc", ".docx", ".xlsx", ".csv", ".txt"],
            "chunk_size": config.ai.chunk_size,
            "chunk_overlap": config.ai.chunk_overlap,
            "timeout_seconds": config.ai.timeout_seconds,
            "max_file_size_mb": config.upload.max_file_size_mb
        }

# Global document processor instance
document_processor = DocumentProcessor()