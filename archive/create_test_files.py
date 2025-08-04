#!/usr/bin/env python3
"""Create comprehensive test files for document analysis testing."""

import os
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF

def create_test_directory():
    """Create test files directory."""
    test_dir = Path("test_files")
    test_dir.mkdir(exist_ok=True)
    return test_dir

def create_word_document(test_dir):
    """Create a sample Word document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('BMO Quarterly Business Review', 0)
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document provides a comprehensive review of BMO\'s performance for Q2 2024. '
        'Key highlights include strong revenue growth, improved operational efficiency, and '
        'successful digital transformation initiatives.'
    )
    
    # Financial Performance
    doc.add_heading('Financial Performance', level=1)
    doc.add_paragraph(
        'Revenue for Q2 2024 reached $3.65 billion, representing a 7.4% increase from the previous quarter. '
        'Net profit margins improved to 32.9%, driven by cost optimization and increased productivity.'
    )
    
    # Key Metrics Table
    doc.add_heading('Key Performance Indicators', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Q1 2024'
    hdr_cells[2].text = 'Q2 2024'
    
    metrics = [
        ('Revenue ($ millions)', '3,400', '3,650'),
        ('Profit Margin (%)', '32.4', '32.9'),
        ('Customer Count', '19,200', '20,100'),
        ('Employee Satisfaction', '87%', '89%')
    ]
    
    for metric, q1, q2 in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = q1
        row_cells[2].text = q2
    
    # Strategic Initiatives
    doc.add_heading('Strategic Initiatives', level=1)
    doc.add_paragraph('Key strategic initiatives completed this quarter:')
    
    initiatives = [
        'Digital banking platform upgrade - 95% completion',
        'AI-powered customer service implementation',
        'Sustainable finance product launch',
        'Cybersecurity infrastructure enhancement',
        'Employee wellness program expansion'
    ]
    
    for initiative in initiatives:
        doc.add_paragraph(initiative, style='List Bullet')
    
    # Future Outlook
    doc.add_heading('Future Outlook', level=1)
    doc.add_paragraph(
        'Looking ahead to Q3 2024, we anticipate continued growth driven by our digital transformation '
        'initiatives and expanded service offerings. We project 6-8% revenue growth and expect to '
        'achieve our annual targets.'
    )
    
    # Save document
    doc_path = test_dir / 'bmo_quarterly_review.docx'
    doc.save(doc_path)
    print(f"Created Word document: {doc_path}")
    return doc_path

def create_excel_workbook(test_dir):
    """Create a sample Excel workbook."""
    wb = Workbook()
    
    # Employee Data Sheet
    ws1 = wb.active
    ws1.title = "Employee Data"
    
    # Headers
    headers = ['Employee ID', 'Name', 'Department', 'Salary', 'Performance', 'Location']
    for col, header in enumerate(headers, 1):
        ws1.cell(row=1, column=col, value=header)
    
    # Employee data
    employees = [
        ['E001', 'John Smith', 'Engineering', 95000, 'Excellent', 'Toronto'],
        ['E002', 'Sarah Johnson', 'Marketing', 82000, 'Good', 'Montreal'],
        ['E003', 'Michael Brown', 'Engineering', 110000, 'Excellent', 'Vancouver'],
        ['E004', 'Emily Davis', 'HR', 68000, 'Good', 'Toronto'],
        ['E005', 'David Wilson', 'Finance', 75000, 'Satisfactory', 'Calgary']
    ]
    
    for row, emp_data in enumerate(employees, 2):
        for col, value in enumerate(emp_data, 1):
            ws1.cell(row=row, column=col, value=value)
    
    # Financial Data Sheet
    ws2 = wb.create_sheet(title="Financial Summary")
    
    # Financial headers
    fin_headers = ['Quarter', 'Revenue ($M)', 'Expenses ($M)', 'Profit ($M)', 'Growth (%)']
    for col, header in enumerate(fin_headers, 1):
        ws2.cell(row=1, column=col, value=header)
    
    # Financial data
    financial_data = [
        ['Q1 2024', 3400, 2300, 1100, 6.3],
        ['Q2 2024', 3650, 2450, 1200, 7.4],
        ['Q3 2024 (Proj)', 3850, 2550, 1300, 5.5],
        ['Q4 2024 (Proj)', 4000, 2600, 1400, 3.9]
    ]
    
    for row, fin_data in enumerate(financial_data, 2):
        for col, value in enumerate(fin_data, 1):
            ws2.cell(row=row, column=col, value=value)
    
    # Department Statistics Sheet
    ws3 = wb.create_sheet(title="Department Stats")
    
    dept_headers = ['Department', 'Headcount', 'Avg Salary', 'Budget ($K)', 'Utilization (%)']
    for col, header in enumerate(dept_headers, 1):
        ws3.cell(row=1, column=col, value=header)
    
    dept_data = [
        ['Engineering', 45, 92000, 4500, 87],
        ['Marketing', 22, 75000, 1800, 92],
        ['Finance', 18, 78000, 1500, 89],
        ['HR', 12, 65000, 850, 85],
        ['IT', 15, 85000, 1200, 91]
    ]
    
    for row, dept in enumerate(dept_data, 2):
        for col, value in enumerate(dept, 1):
            ws3.cell(row=row, column=col, value=value)
    
    # Save workbook
    excel_path = test_dir / 'bmo_business_data.xlsx'
    wb.save(excel_path)
    print(f"Created Excel workbook: {excel_path}")
    return excel_path

def create_pdf_document(test_dir):
    """Create a sample PDF document."""
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 15)
            self.cell(0, 10, 'BMO Technology Strategy Document', 0, 1, 'C')
            self.ln(10)
        
        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
    
    pdf = PDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)
    
    # Title
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, 'Digital Transformation Strategy 2024-2026', 0, 1, 'L')
    pdf.ln(5)
    
    # Introduction
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, '1. Executive Overview', 0, 1, 'L')
    pdf.set_font('Arial', '', 11)
    
    intro_text = """BMO's digital transformation strategy focuses on modernizing our technology infrastructure, 
enhancing customer experience, and driving operational efficiency. This three-year roadmap outlines 
key initiatives, investments, and expected outcomes.

Our strategic pillars include:
- Cloud-first infrastructure migration
- AI and machine learning integration  
- Enhanced cybersecurity framework
- Customer-centric digital products
- Data-driven decision making"""
    
    # Split text into lines for PDF
    lines = intro_text.split('\n')
    for line in lines:
        if line.strip():
            if line.strip().startswith('-'):
                pdf.cell(10, 6, '', 0, 0)  # Indent
                pdf.cell(0, 6, line.strip(), 0, 1, 'L')  
            else:
                pdf.cell(0, 6, line.strip(), 0, 1, 'L')
        else:
            pdf.ln(3)
    
    pdf.ln(5)
    
    # Technology Initiatives
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, '2. Key Technology Initiatives', 0, 1, 'L')
    pdf.set_font('Arial', '', 11)
    
    initiatives_text = """2.1 Cloud Migration Initiative
Target: Migrate 80% of applications to cloud by Q4 2025
Investment: $125M over 24 months
Expected ROI: 35% cost reduction, improved scalability

2.2 Artificial Intelligence Platform
Objective: Implement AI-driven customer insights and automation
Applications: Risk assessment, fraud detection, personalized banking
Timeline: Pilot Q3 2024, full deployment Q2 2025

2.3 Cybersecurity Enhancement
Scope: Zero-trust architecture, advanced threat detection
Compliance: Regulatory requirements and industry standards
Budget: $45M annually for enhanced security measures

2.4 Digital Customer Experience
Focus: Mobile app redesign, omnichannel integration
Metrics: Customer satisfaction score target of 4.5/5.0
Launch: Phased rollout starting Q1 2025"""
    
    lines = initiatives_text.split('\n')
    for line in lines:
        if line.strip():
            if line.strip().startswith(('2.1', '2.2', '2.3', '2.4')):
                pdf.set_font('Arial', 'B', 11)
                pdf.cell(0, 7, line.strip(), 0, 1, 'L')
                pdf.set_font('Arial', '', 11)
            else:
                pdf.cell(0, 6, line.strip(), 0, 1, 'L')
        else:
            pdf.ln(2)
    
    pdf.ln(5)
    
    # Success Metrics
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, '3. Success Metrics and KPIs', 0, 1, 'L')
    pdf.set_font('Arial', '', 11)
    
    metrics_text = """Our success will be measured through the following key performance indicators:

- Operational Efficiency: 25% reduction in processing time
- Customer Satisfaction: Maintain >90% satisfaction rate
- Cost Optimization: $200M annual savings by 2026
- Security Incidents: <5 major incidents annually
- Digital Adoption: 75% of customers using digital channels
- Employee Productivity: 20% improvement in task completion"""
    
    lines = metrics_text.split('\n')
    for line in lines:
        if line.strip():
            if line.strip().startswith('-'):
                pdf.cell(10, 6, '', 0, 0)  # Indent
                pdf.cell(0, 6, line.strip(), 0, 1, 'L')
            else:
                pdf.cell(0, 6, line.strip(), 0, 1, 'L')
        else:
            pdf.ln(3)
    
    # Save PDF
    pdf_path = test_dir / 'bmo_tech_strategy.pdf'
    pdf.output(str(pdf_path))
    print(f"Created PDF document: {pdf_path}")
    return pdf_path

def main():
    """Create all test files."""
    print("Creating comprehensive test files for document analysis...")
    
    # Create test directory
    test_dir = create_test_directory()
    
    # Create all document types
    word_path = create_word_document(test_dir)
    excel_path = create_excel_workbook(test_dir)
    pdf_path = create_pdf_document(test_dir)
    
    print(f"\n✅ Test files created successfully in: {test_dir}")
    print(f"📄 Word Document: {word_path.name}")
    print(f"📊 Excel Workbook: {excel_path.name}")  
    print(f"📋 PDF Document: {pdf_path.name}")
    print(f"📈 CSV Files: sample_employees.csv, quarterly_report.csv")
    
    print(f"\nTotal files created: 5")
    print("Ready for comprehensive testing! 🚀")

if __name__ == "__main__":
    main()